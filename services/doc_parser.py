"""Document parser for Word (.docx), PDF, Markdown, and Feishu docs."""

import logging
import os
import re
import time

import requests

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Feishu API config — loaded from environment or config.json
# ---------------------------------------------------------------------------
FEISHU_APP_ID = os.environ.get("FEISHU_APP_ID", "")
FEISHU_APP_SECRET = os.environ.get("FEISHU_APP_SECRET", "")
FEISHU_DOMAIN = os.environ.get("FEISHU_DOMAIN", "https://open.feishu.cn")

_tenant_token_cache = {"token": "", "expires": 0}

# Regex to match Feishu docx URLs
FEISHU_URL_PATTERN = re.compile(
    r"https?://[a-zA-Z0-9\-]+\.feishu\.cn/(?:docx|wiki)/([a-zA-Z0-9]+)"
)


def is_feishu_url(text: str) -> bool:
    """Check if the given text is a Feishu document URL."""
    return bool(FEISHU_URL_PATTERN.search(text))


def extract_feishu_token(url: str) -> str:
    """Extract document token from a Feishu URL."""
    m = FEISHU_URL_PATTERN.search(url)
    if not m:
        raise ValueError(f"Cannot extract doc token from URL: {url}")
    return m.group(1)


def parse_document(file_path: str) -> str:
    """Parse a document and return its text content.

    Supports .docx, .pdf, .md, .txt files, and Feishu document URLs.
    """
    # Check if it's a Feishu URL instead of a file path
    if isinstance(file_path, str) and is_feishu_url(file_path):
        logger.info("Parsing Feishu document: %s", file_path)
        text = _parse_feishu(file_path)
        if not text.strip():
            raise ValueError("Feishu document appears to be empty after parsing.")
        return text

    ext = os.path.splitext(file_path)[1].lower()
    parsers = {
        ".docx": _parse_docx,
        ".pdf": _parse_pdf,
        ".md": _parse_markdown,
        ".markdown": _parse_markdown,
        ".txt": _parse_plaintext,
    }
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported document format: {ext}")

    logger.info("Parsing document: %s", file_path)
    text = parser(file_path)
    if not text.strip():
        raise ValueError("Document appears to be empty after parsing.")
    return text


# ---------------------------------------------------------------------------
# Feishu document parsing
# ---------------------------------------------------------------------------

def get_feishu_config():
    """Load Feishu config, first from env then from config.json."""
    from config import load_config
    app_id = FEISHU_APP_ID
    app_secret = FEISHU_APP_SECRET
    domain = FEISHU_DOMAIN

    if not app_id or not app_secret:
        cfg = load_config()
        app_id = cfg.get("feishu_app_id", "")
        app_secret = cfg.get("feishu_app_secret", "")
        domain = cfg.get("feishu_domain", domain)

    if not app_id or not app_secret:
        raise ValueError(
            "Feishu App ID and App Secret are required. "
            "Configure them in the Settings page or set FEISHU_APP_ID / FEISHU_APP_SECRET env vars."
        )
    return app_id, app_secret, domain


def get_tenant_access_token() -> str:
    """Get Feishu tenant_access_token with caching."""
    now = time.time()
    if _tenant_token_cache["token"] and _tenant_token_cache["expires"] > now + 60:
        return _tenant_token_cache["token"]

    app_id, app_secret, domain = get_feishu_config()
    url = f"{domain}/open-apis/auth/v3/tenant_access_token/internal"
    resp = requests.post(url, json={"app_id": app_id, "app_secret": app_secret}, timeout=10)
    resp.raise_for_status()
    data = resp.json()

    if data.get("code") != 0:
        raise ValueError(f"Failed to get Feishu token: {data.get('msg', 'unknown error')}")

    token = data["tenant_access_token"]
    expire = data.get("expire", 7200)
    _tenant_token_cache["token"] = token
    _tenant_token_cache["expires"] = now + expire

    logger.info("Feishu tenant_access_token refreshed, expires in %ds", expire)
    return token


def _feishu_api_get(path: str, params: dict = None) -> dict:
    """Make an authenticated GET request to Feishu Open API."""
    _, _, domain = get_feishu_config()
    token = get_tenant_access_token()
    url = f"{domain}/open-apis{path}"
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(url, headers=headers, params=params, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    if data.get("code") != 0:
        raise ValueError(f"Feishu API error: {data.get('msg', 'unknown')} (code={data.get('code')})")
    return data.get("data", {})


def _parse_feishu(url: str) -> str:
    """Parse a Feishu document by URL, returning structured text."""
    doc_token = extract_feishu_token(url)

    # Check if this is a wiki URL — need to resolve to actual doc token
    if "/wiki/" in url:
        doc_token = _resolve_wiki_token(doc_token)

    # Step 1: Get document basic info
    try:
        doc_info = _feishu_api_get(f"/docx/v1/documents/{doc_token}")
        title = doc_info.get("document", {}).get("title", "Untitled")
    except Exception:
        title = "Untitled"

    # Step 2: Get all blocks for full structured content
    blocks = _get_all_blocks(doc_token)

    # Step 3: Convert blocks to readable text
    text = _blocks_to_text(blocks, title)
    return text


def _resolve_wiki_token(wiki_token: str) -> str:
    """Resolve a wiki node token to the actual document token."""
    try:
        data = _feishu_api_get(f"/wiki/v2/spaces/get_node", params={"token": wiki_token})
        node = data.get("node", {})
        obj_token = node.get("obj_token", wiki_token)
        logger.info("Resolved wiki token %s -> doc token %s", wiki_token, obj_token)
        return obj_token
    except Exception as e:
        logger.warning("Failed to resolve wiki token, using as-is: %s", e)
        return wiki_token


def _get_all_blocks(doc_token: str) -> list:
    """Fetch all blocks from a Feishu document with pagination."""
    all_blocks = []
    page_token = None

    while True:
        params = {"document_id": doc_token, "page_size": 500}
        if page_token:
            params["page_token"] = page_token

        data = _feishu_api_get(f"/docx/v1/documents/{doc_token}/blocks", params=params)
        items = data.get("items", [])
        all_blocks.extend(items)

        if not data.get("has_more"):
            break
        page_token = data.get("page_token")

    return all_blocks


def _extract_text_from_elements(elements: list) -> str:
    """Extract plain text from Feishu block text elements."""
    parts = []
    for el in elements:
        if "text_run" in el:
            parts.append(el["text_run"].get("content", ""))
        elif "mention_user" in el:
            parts.append(f"@{el['mention_user'].get('user_id', 'user')}")
        elif "equation" in el:
            parts.append(el["equation"].get("content", ""))
    return "".join(parts)


def _blocks_to_text(blocks: list, title: str = "") -> str:
    """Convert Feishu document blocks into readable structured text."""
    parts = []
    if title:
        parts.append(f"# {title}\n")

    # Build a map for nested blocks
    block_map = {b.get("block_id"): b for b in blocks}

    for block in blocks:
        block_type = block.get("block_type")

        # Page block (root) — skip
        if block_type == 1:
            continue

        # Text / Paragraph
        if block_type == 2:
            text_body = block.get("text", {})
            elements = text_body.get("elements", [])
            text = _extract_text_from_elements(elements)
            if text.strip():
                parts.append(text)

        # Heading 1-9
        elif block_type in range(3, 12):
            level = block_type - 2  # 3->H1, 4->H2, etc.
            heading = block.get(f"heading{level}", block.get("heading1", block.get("heading2", block.get("heading3", {}))))
            # Try all possible heading keys
            for i in range(1, 10):
                key = f"heading{i}"
                if key in block:
                    heading = block[key]
                    break
            elements = heading.get("elements", [])
            text = _extract_text_from_elements(elements)
            if text.strip():
                parts.append(f"{'#' * level} {text}")

        # Bullet list
        elif block_type == 12:
            bullet = block.get("bullet", {})
            elements = bullet.get("elements", [])
            text = _extract_text_from_elements(elements)
            if text.strip():
                parts.append(f"- {text}")

        # Ordered list
        elif block_type == 13:
            ordered = block.get("ordered", {})
            elements = ordered.get("elements", [])
            text = _extract_text_from_elements(elements)
            if text.strip():
                parts.append(f"1. {text}")

        # Code block
        elif block_type == 14:
            code = block.get("code", {})
            elements = code.get("elements", [])
            text = _extract_text_from_elements(elements)
            lang = code.get("style", {}).get("language", "")
            if text.strip():
                parts.append(f"```{lang}\n{text}\n```")

        # Quote
        elif block_type == 15:
            quote = block.get("quote", {})
            elements = quote.get("elements", [])
            text = _extract_text_from_elements(elements)
            if text.strip():
                parts.append(f"> {text}")

        # Todo / Checkbox
        elif block_type == 17:
            todo = block.get("todo", {})
            elements = todo.get("elements", [])
            text = _extract_text_from_elements(elements)
            done = todo.get("style", {}).get("done", False)
            mark = "x" if done else " "
            if text.strip():
                parts.append(f"- [{mark}] {text}")

        # Divider
        elif block_type == 22:
            parts.append("---")

        # Table (block_type 31) — extract cell content
        elif block_type == 31:
            table_info = block.get("table", {})
            row_size = table_info.get("property", {}).get("row_size", 0)
            col_size = table_info.get("property", {}).get("column_size", 0)
            cells = table_info.get("cells", [])

            if cells and row_size and col_size:
                table_rows = []
                for r in range(row_size):
                    row_cells = []
                    for c in range(col_size):
                        idx = r * col_size + c
                        if idx < len(cells):
                            cell_id = cells[idx]
                            cell_block = block_map.get(cell_id, {})
                            # Cell blocks contain child blocks; extract text
                            child_ids = cell_block.get("children", [])
                            cell_texts = []
                            for cid in child_ids:
                                cb = block_map.get(cid, {})
                                if cb.get("block_type") == 2:
                                    els = cb.get("text", {}).get("elements", [])
                                    cell_texts.append(_extract_text_from_elements(els))
                            row_cells.append(" ".join(cell_texts))
                        else:
                            row_cells.append("")
                    table_rows.append(row_cells)

                if table_rows:
                    # Format as markdown-ish table
                    header = " | ".join(table_rows[0])
                    separator = " | ".join(["---"] * col_size)
                    lines = [header, separator]
                    for row in table_rows[1:]:
                        lines.append(" | ".join(row))
                    parts.append("\n".join(lines))

        # Callout
        elif block_type == 19:
            callout = block.get("callout", {})
            # Callout contains child blocks
            child_ids = block.get("children", [])
            for cid in child_ids:
                cb = block_map.get(cid, {})
                if cb.get("block_type") == 2:
                    els = cb.get("text", {}).get("elements", [])
                    text = _extract_text_from_elements(els)
                    if text.strip():
                        parts.append(f"> {text}")

    return "\n\n".join(parts)


def _parse_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading structure
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            try:
                level = int(level)
            except ValueError:
                level = 1
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _parse_pdf(file_path: str) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        pass

    # Fallback to PyPDF2
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _parse_markdown(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _parse_plaintext(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()
